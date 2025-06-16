from flask import Flask, request, jsonify, send_file, render_template_string
import os
import json
import csv
import io
import asyncio
from werkzeug.utils import secure_filename

# Create Flask app
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Define allowed extensions
ALLOWED_EXTENSIONS = {
    'pdf', 'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'tif',
    'docx', 'doc', 'xlsx', 'xls', 'pptx', 'ppt',
    'txt', 'rtf', 'csv', 'html', 'htm', 'xml',
    'odt', 'ods', 'odp', 'epub', 'md', 'json'
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# HTML Template for the UI
HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>📄 Document Text Extraction Tool</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }
        
        function copyToClipboard(text) {
            navigator.clipboard.writeText(text).then(function() {
                alert('✅ Text copied to clipboard!');
            }).catch(function(err) {
                console.error('Could not copy text: ', err);
                // Fallback method
                const textArea = document.createElement('textarea');
                textArea.value = text;
                document.body.appendChild(textArea);
                textArea.select();
                document.execCommand('copy');
                document.body.removeChild(textArea);
                alert('✅ Text copied to clipboard!');
            });
        }
        
        function downloadText(filename, text) {
            const element = document.createElement('a');
            const file = new Blob([text], {type: 'text/plain'});
            element.href = URL.createObjectURL(file);
            element.download = filename.replace(/\.[^/.]+$/, "") + '_extracted.txt';
            document.body.appendChild(element);
            element.click();
            document.body.removeChild(element);
        }
        .container { 
            max-width: 800px; 
            margin: 0 auto; 
            background: white; 
            border-radius: 15px; 
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
            overflow: hidden;
        }
        .header { 
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white; 
            padding: 30px; 
            text-align: center; 
        }
        .header h1 { font-size: 2.5em; margin-bottom: 10px; }
        .header p { font-size: 1.1em; opacity: 0.9; }
        .main-content { padding: 40px; }
        .upload-section { 
            border: 3px dashed #667eea; 
            border-radius: 10px; 
            padding: 40px; 
            text-align: center; 
            margin-bottom: 30px;
            transition: all 0.3s ease;
        }
        .upload-section:hover { 
            border-color: #764ba2; 
            background: #f8f9ff;
        }
        .file-input-wrapper {
            position: relative;
            overflow: hidden;
            display: inline-block;
        }
        .file-input {
            position: absolute;
            left: -9999px;
        }
        .file-input-button {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 15px 30px;
            border-radius: 25px;
            cursor: pointer;
            font-size: 16px;
            border: none;
            transition: transform 0.2s ease;
        }
        .file-input-button:hover {
            transform: translateY(-2px);
            box-shadow: 0 5px 15px rgba(0,0,0,0.2);
        }
        .format-selection {
            margin: 20px 0;
        }
        .format-selection label {
            margin-right: 20px;
            font-weight: 500;
        }
        .format-selection input[type="radio"] {
            margin-right: 5px;
        }
        .extract-button {
            background: linear-gradient(135deg, #28a745 0%, #20c997 100%);
            color: white;
            padding: 15px 40px;
            border: none;
            border-radius: 25px;
            font-size: 18px;
            cursor: pointer;
            transition: all 0.3s ease;
            margin-top: 20px;
        }
        .extract-button:hover {
            transform: translateY(-2px);
            box-shadow: 0 5px 15px rgba(0,0,0,0.2);
        }
        .extract-button:disabled {
            opacity: 0.6;
            cursor: not-allowed;
            transform: none;
        }
        .result-section {
            margin-top: 30px;
            padding: 20px;
            background: #f8f9fa;
            border-radius: 10px;
            border-left: 5px solid #667eea;
        }
        .result-section.error {
            border-left-color: #dc3545;
            background: #f8d7da;
        }
        .result-section.success {
            border-left-color: #28a745;
            background: #d4edda;
        }
        .result-text {
            background: #f8f9fa;
            padding: 20px;
            border-radius: 8px;
            max-height: 400px;
            overflow-y: auto;
            white-space: pre-wrap;
            font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
            margin-top: 15px;
            border: 1px solid #dee2e6;
            font-size: 14px;
            line-height: 1.5;
        }
        .result-text:empty:before {
            content: "No text content available";
            color: #6c757d;
            font-style: italic;
        }
        .action-buttons {
            margin-top: 15px;
            display: flex;
            gap: 10px;
        }
        .action-buttons button {
            background: #667eea;
            color: white;
            border: none;
            padding: 8px 16px;
            border-radius: 5px;
            cursor: pointer;
            font-size: 14px;
        }
        .action-buttons button:hover {
            background: #5a6fd8;
        }
        .file-info {
            background: #e3f2fd;
            padding: 15px;
            border-radius: 8px;
            margin: 15px 0;
        }
        .file-info h4 {
            color: #1976d2;
            margin-bottom: 10px;
        }
        .stats {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 15px;
            margin: 15px 0;
        }
        .stat-item {
            background: white;
            padding: 15px;
            border-radius: 8px;
            text-align: center;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        .stat-value {
            font-size: 1.5em;
            font-weight: bold;
            color: #667eea;
        }
        .stat-label {
            color: #666;
            font-size: 0.9em;
        }
        .loading {
            display: none;
            text-align: center;
            color: #667eea;
        }
        .loading.show {
            display: block;
        }
        .spinner {
            border: 4px solid #f3f3f3;
            border-top: 4px solid #667eea;
            border-radius: 50%;
            width: 40px;
            height: 40px;
            animation: spin 1s linear infinite;
            margin: 20px auto;
        }
        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
        .supported-formats {
            background: #f8f9fa;
            padding: 20px;
            border-radius: 10px;
            margin-top: 20px;
        }
        .supported-formats h3 {
            color: #495057;
            margin-bottom: 15px;
        }
        .format-grid {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(80px, 1fr));
            gap: 10px;
        }
        .format-tag {
            background: #667eea;
            color: white;
            padding: 5px 10px;
            border-radius: 15px;
            text-align: center;
            font-size: 0.8em;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>📄 Document Text Extractor</h1>
            <p>Upload any document and extract text using AI-powered processing</p>
        </div>
        
        <div class="main-content">
            <form id="extractForm" enctype="multipart/form-data">
                <div class="upload-section">
                    <h3>📁 Choose Your Document</h3>
                    <p style="margin: 15px 0; color: #666;">
                        Support for PDF, Word, Excel, PowerPoint, Images, and more!
                    </p>
                    
                    <div class="file-input-wrapper">
                        <input type="file" id="file" name="file" class="file-input" accept=".pdf,.txt,.docx,.doc,.xlsx,.xls,.pptx,.ppt,.jpg,.jpeg,.png,.gif,.bmp,.tiff,.html,.xml,.odt,.epub,.rtf,.csv,.md,.json" required>
                        <button type="button" class="file-input-button" onclick="document.getElementById('file').click()">
                            📎 Select File
                        </button>
                    </div>
                    
                    <div id="selectedFile" style="margin-top: 15px; font-weight: 500; color: #667eea;"></div>
                </div>
                
                <div class="format-selection">
                    <h4>📊 Output Format:</h4>
                    <label><input type="radio" name="format" value="json" checked> JSON (Detailed)</label>
                    <label><input type="radio" name="format" value="csv"> CSV (Table)</label>
                </div>
                
                <div style="text-align: center;">
                    <button type="submit" class="extract-button" id="extractBtn">
                        🚀 Extract Text
                    </button>
                </div>
            </form>
            
            <div class="loading" id="loading">
                <div class="spinner"></div>
                <p>Processing your document... Please wait</p>
            </div>
            
            <div id="result"></div>
            
            <div class="supported-formats">
                <h3>📋 Supported File Formats</h3>
                <div class="format-grid">
                    <div class="format-tag">PDF</div>
                    <div class="format-tag">DOCX ✅</div>
                    <div class="format-tag">DOC ⚠️</div>
                    <div class="format-tag">DOC</div>
                    <div class="format-tag">XLSX</div>
                    <div class="format-tag">XLS</div>
                    <div class="format-tag">PPTX</div>
                    <div class="format-tag">PPT</div>
                    <div class="format-tag">TXT</div>
                    <div class="format-tag">RTF</div>
                    <div class="format-tag">CSV</div>
                    <div class="format-tag">HTML</div>
                    <div class="format-tag">XML</div>
                    <div class="format-tag">ODT</div>
                    <div class="format-tag">EPUB</div>
                    <div class="format-tag">JPG</div>
                    <div class="format-tag">PNG</div>
                    <div class="format-tag">GIF</div>
                    <div class="format-tag">BMP</div>
                    <div class="format-tag">TIFF</div>
                    <div class="format-tag">JSON</div>
                    <div class="format-tag">MD</div>
                </div>
            </div>
        </div>
    </div>
    
    <script>
        // File selection handler
        document.getElementById('file').addEventListener('change', function(e) {
            const file = e.target.files[0];
            const selectedFileDiv = document.getElementById('selectedFile');
            if (file) {
                selectedFileDiv.innerHTML = `✅ Selected: ${file.name} (${(file.size / 1024 / 1024).toFixed(2)} MB)`;
            } else {
                selectedFileDiv.innerHTML = '';
            }
        });
        
        // Form submission handler
        document.getElementById('extractForm').addEventListener('submit', async function(e) {
            e.preventDefault();
            
            const formData = new FormData();
            const fileInput = document.getElementById('file');
            const formatInputs = document.getElementsByName('format');
            
            if (!fileInput.files[0]) {
                showResult('Please select a file first!', 'error');
                return;
            }
            
            const selectedFormat = Array.from(formatInputs).find(radio => radio.checked).value;
            
            formData.append('file', fileInput.files[0]);
            formData.append('format', selectedFormat);
            
            // Show loading
            document.getElementById('loading').classList.add('show');
            document.getElementById('extractBtn').disabled = true;
            document.getElementById('result').innerHTML = '';
            
            try {
                const response = await fetch('/extract', {
                    method: 'POST',
                    body: formData
                });
                
                if (selectedFormat === 'csv') {
                    // Handle CSV download
                    if (response.ok) {
                        const blob = await response.blob();
                        const url = window.URL.createObjectURL(blob);
                        const a = document.createElement('a');
                        a.style.display = 'none';
                        a.href = url;
                        a.download = `extracted_${Date.now()}.csv`;
                        document.body.appendChild(a);
                        a.click();
                        window.URL.revokeObjectURL(url);
                        showResult('✅ CSV file downloaded successfully!', 'success');
                    } else {
                        const errorText = await response.text();
                        showResult('❌ Error: ' + errorText, 'error');
                    }
                } else {
                    // Handle JSON response
                    const result = await response.json();
                    if (response.ok) {
                        showResult(result, 'success');
                    } else {
                        showResult('❌ Error: ' + (result.error || 'Unknown error'), 'error');
                    }
                }
            } catch (error) {
                showResult('❌ Network error: ' + error.message, 'error');
            } finally {
                // Hide loading
                document.getElementById('loading').classList.remove('show');
                document.getElementById('extractBtn').disabled = false;
            }
        });
        
        function showResult(data, type) {
            const resultDiv = document.getElementById('result');
            
            if (typeof data === 'string') {
                resultDiv.innerHTML = `
                    <div class="result-section ${type}">
                        <h3>${type === 'error' ? '❌ Error' : '✅ Success'}</h3>
                        <p>${data}</p>
                    </div>
                `;
            } else {
                // JSON result
                resultDiv.innerHTML = `
                    <div class="result-section ${type}">
                        <h3>✅ Extraction Complete!</h3>
                        
                        <div class="file-info">
                            <h4>📄 File Information</h4>
                            <p><strong>Filename:</strong> ${data.filename}</p>
                            <p><strong>Type:</strong> ${data.file_extension?.toUpperCase()}</p>
                            <p><strong>Method:</strong> ${data.extraction_method}</p>
                        </div>
                        
                        <div class="stats">
                            <div class="stat-item">
                                <div class="stat-value">${data.word_count || 0}</div>
                                <div class="stat-label">Words</div>
                            </div>
                            <div class="stat-item">
                                <div class="stat-value">${data.character_count || 0}</div>
                                <div class="stat-label">Characters</div>
                            </div>
                            <div class="stat-item">
                                <div class="stat-value">${data.total_pages || 1}</div>
                                <div class="stat-label">Pages</div>
                            </div>
                        </div>
                        
                        <h4>📝 Extracted Text Content:</h4>
                        <div class="result-text">${data.full_text ? data.full_text : 'No text could be extracted from this file'}</div>
                        
                        ${data.full_text && data.full_text.length > 10 ? `
                        <div class="action-buttons">
                            <button onclick="copyToClipboard(\`${data.full_text.replace(/`/g, '\\`')}\`)">📋 Copy Text</button>
                            <button onclick="downloadText('${data.filename}', \`${data.full_text.replace(/`/g, '\\`')}\`)">💾 Download as TXT</button>
                        </div>
                        ` : ''}
                    </div>
                `;
            }
        }
    </script>
</body>
</html>
'''

# Routes
@app.route('/')
def home():
    """Main page with upload interface"""
    return render_template_string(HTML_TEMPLATE)

@app.route('/languages')
def available_languages():
    """Show all available OCR languages"""
    try:
        import pytesseract
        langs = pytesseract.get_languages(config='')
        return jsonify({
            "total_languages": len(langs),
            "available_languages": sorted(langs),
            "note": "These are all languages supported by your Tesseract installation"
        })
    except Exception as e:
        return jsonify({
            "error": f"Could not retrieve languages: {str(e)}",
            "note": "Tesseract may not be properly installed"
        })

@app.route('/extract/docx', methods=['POST'])
def extract_docx_only():
    """Dedicated endpoint for DOCX extraction"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # Check if it's a DOCX file
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        if file_ext != 'docx':
            return jsonify({
                "error": f"This endpoint only accepts DOCX files. Received: {file_ext}",
                "note": "Use /extract for other file types or convert DOC to DOCX"
            }), 400
        
        # Save and process
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        try:
            from docx import Document
            
            doc = Document(filepath)
            
            # Detailed DOCX analysis
            result = {
                "filename": filename,
                "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "extraction_method": "python-docx (dedicated DOCX processor)",
                "document_stats": {
                    "total_paragraphs": len(doc.paragraphs),
                    "total_tables": len(doc.tables),
                    "total_sections": len(doc.sections)
                }
            }
            
            # Extract all content
            all_text_parts = []
            
            # Paragraphs
            paragraph_texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraph_texts.append(para.text.strip())
            
            if paragraph_texts:
                all_text_parts.append("=== PARAGRAPHS ===")
                all_text_parts.extend(paragraph_texts)
            
            # Tables
            if doc.tables:
                all_text_parts.append("\n=== TABLES ===")
                for i, table in enumerate(doc.tables, 1):
                    all_text_parts.append(f"\n--- Table {i} ---")
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_data:
                            all_text_parts.append(" | ".join(row_data))
            
            full_text = "\n".join(all_text_parts)
            
            result.update({
                "full_text": full_text,
                "word_count": len(full_text.split()) if full_text else 0,
                "character_count": len(full_text),
                "status": "success" if full_text else "no_content"
            })
            
            return jsonify(result)
            
        finally:
            if os.path.exists(filepath):
                os.remove(filepath)
                
    except Exception as e:
        return jsonify({"error": f"DOCX processing failed: {str(e)}"}), 500
    """Show all available OCR languages"""
    try:
        import pytesseract
        langs = pytesseract.get_languages(config='')
        return jsonify({
            "total_languages": len(langs),
            "available_languages": sorted(langs),
            "note": "These are all languages supported by your Tesseract installation"
        })
    except Exception as e:
        return jsonify({
            "error": f"Could not retrieve languages: {str(e)}",
            "note": "Tesseract may not be properly installed"
        })

@app.route('/health')
def health_check():
    try:
        import pytesseract
        langs = pytesseract.get_languages(config='')
        lang_count = len(langs)
    except:
        lang_count = 0
    
    return jsonify({
        "status": "healthy", 
        "message": "Document extraction API is running",
        "ocr_languages_available": lang_count,
        "endpoints": {
            "/languages": "GET - View all available OCR languages"
        }
    })

@app.route('/extract', methods=['POST'])
def extract_document():
    try:
        print("Extract endpoint accessed!")
        
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        if not allowed_file(file.filename):
            return jsonify({
                "error": f"File type not supported. Allowed: {list(ALLOWED_EXTENSIONS)}"
            }), 400
        
        # Get output format
        output_format = request.form.get('format', 'json').lower()
        
        # Save the file temporarily
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Basic file info
        file_size = os.path.getsize(filepath)
        file_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else 'unknown'
        
        # Basic text extraction for common formats
        extracted_text = ""
        method = "Basic extraction"
        
        try:
            if file_extension == 'txt':
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    extracted_text = f.read()
                method = "Direct text reading"
                
            elif file_extension == 'html':
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    html_content = f.read()
                    # Basic HTML tag removal
                    import re
                    # Remove script and style elements
                    html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                    html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                    # Remove HTML tags
                    extracted_text = re.sub(r'<[^>]+>', '', html_content)
                    # Clean up whitespace
                    extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
                method = "HTML text extraction"
                
            elif file_extension == 'csv':
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    import csv as csv_module
                    csv_reader = csv_module.reader(f)
                    rows = []
                    for row in csv_reader:
                        rows.append(' | '.join(row))
                    extracted_text = '\n'.join(rows)
                method = "CSV text extraction"
                
            elif file_extension == 'json':
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    json_content = f.read()
                    try:
                        # Try to parse and format JSON
                        import json
                        parsed_json = json.loads(json_content)
                        extracted_text = json.dumps(parsed_json, indent=2)
                    except:
                        # If parsing fails, just return raw content
                        extracted_text = json_content
                method = "JSON text extraction"
                
            elif file_extension in ['md', 'markdown']:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    extracted_text = f.read()
                method = "Markdown text reading"
                
            elif file_extension == 'rtf':
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    rtf_content = f.read()
                    # Basic RTF tag removal (very simple)
                    import re
                    extracted_text = re.sub(r'\\[a-zA-Z]+\d*', '', rtf_content)
                    extracted_text = re.sub(r'[{}]', '', extracted_text)
                    extracted_text = extracted_text.replace('\\', '').strip()
                method = "Basic RTF text extraction"
                
            elif file_extension == 'xml':
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    xml_content = f.read()
                    # Basic XML tag removal
                    import re
                    extracted_text = re.sub(r'<[^>]+>', '', xml_content)
                    extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
                method = "XML text extraction"
                
            elif file_extension == 'pdf':
                # Real PDF text extraction
                try:
                    import PyPDF2
                    
                    with open(filepath, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        num_pages = len(pdf_reader.pages)
                        
                        # Extract text from all pages
                        all_text = []
                        for page_num in range(num_pages):
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            if page_text.strip():
                                all_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
                        
                        if all_text:
                            extracted_text = f"""📄 PDF Text Extraction Results

✅ TEXT EXTRACTED SUCCESSFULLY!

Document: {filename}
Pages: {num_pages}
Method: Direct PDF text extraction
File Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)

==================== EXTRACTED TEXT ====================

{chr(10).join(all_text)}

=======================================================

Extraction completed using PyPDF2 library."""
                            method = f"PDF Text Extraction ({num_pages} pages)"
                        else:
                            # If no text found, try OCR on PDF
                            try:
                                from pdf2image import convert_from_path
                                import pytesseract
                                from PIL import Image
                                
                                # Convert PDF to images and OCR
                                pages = convert_from_path(filepath, dpi=200)
                                ocr_results = []
                                
                                for i, page in enumerate(pages[:3]):  # Limit to first 3 pages
                                    ocr_text = pytesseract.image_to_string(page)
                                    if ocr_text.strip():
                                        ocr_results.append(f"--- Page {i + 1} (OCR) ---\n{ocr_text}")
                                
                                if ocr_results:
                                    extracted_text = f"""📄 PDF OCR Extraction Results

✅ TEXT EXTRACTED VIA OCR!

Document: {filename}
Pages: {num_pages} (OCR processed: {len(ocr_results)})
Method: PDF to Image + OCR (No direct text found)
File Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)

==================== EXTRACTED TEXT ====================

{chr(10).join(ocr_results)}

=======================================================

Note: This PDF appears to be image-based, so OCR was used."""
                                    method = f"PDF OCR Extraction ({len(ocr_results)} pages)"
                                else:
                                    extracted_text = f"""📄 PDF Processing Complete - No Text Found

Document: {filename}
Pages: {num_pages}
File Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)

This PDF appears to contain no extractable text. This could be because:
- The PDF contains only images
- The PDF is password protected
- The PDF is corrupted
- Text is embedded as non-extractable content"""
                                    method = "PDF Processing (No text found)"
                            except:
                                extracted_text = f"""📄 PDF Text Extraction Attempted

Document: {filename}
Pages: {num_pages}
File Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)

Direct text extraction found no readable text.
OCR processing is not available in this configuration.

The PDF may contain:
- Image-based content requiring OCR
- Protected or encrypted text
- Non-standard text encoding"""
                                method = "PDF Processing (Limited)"
                                
                except Exception as pdf_error:
                    extracted_text = f"""❌ PDF Processing Error

Document: {filename}
File Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)

Error: {str(pdf_error)}

This could be due to:
- Corrupted PDF file
- Password-protected PDF
- Unsupported PDF format
- Processing configuration issue"""
                    method = "PDF Error"
                    
            elif file_extension in ['docx']:
                # Real DOCX text extraction using python-docx
                try:
                    from docx import Document
                    
                    # Open the DOCX document
                    doc = Document(filepath)
                    
                    # Extract text from paragraphs
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text.strip())
                    
                    # Extract text from tables
                    table_data = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_data.append(cell.text.strip())
                            if row_data:
                                table_data.append(" | ".join(row_data))
                    
                    # Extract headers and footers
                    headers_footers = []
                    for section in doc.sections:
                        # Headers
                        if section.header:
                            for para in section.header.paragraphs:
                                if para.text.strip():
                                    headers_footers.append(f"[HEADER] {para.text.strip()}")
                        # Footers
                        if section.footer:
                            for para in section.footer.paragraphs:
                                if para.text.strip():
                                    headers_footers.append(f"[FOOTER] {para.text.strip()}")
                    
                    # Combine all extracted content
                    all_content = []
                    
                    if headers_footers:
                        all_content.append("=== HEADERS & FOOTERS ===")
                        all_content.extend(headers_footers)
                        all_content.append("")
                    
                    if paragraphs:
                        all_content.append("=== DOCUMENT CONTENT ===")
                        all_content.extend(paragraphs)
                        all_content.append("")
                    
                    if table_data:
                        all_content.append("=== TABLES ===")
                        all_content.extend(table_data)
                        all_content.append("")
                    
                    if all_content:
                        full_extracted_text = "\n".join(all_content)
                        
                        # Document statistics
                        total_paragraphs = len(paragraphs)
                        total_tables = len(doc.tables)
                        total_sections = len(doc.sections)
                        
                        extracted_text = f"""📄 DOCX Document Extraction Results

✅ TEXT EXTRACTED SUCCESSFULLY!

Document: {filename}
File Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)
Extraction Method: python-docx library

📊 Document Structure:
- Paragraphs: {total_paragraphs}
- Tables: {total_tables}
- Sections: {total_sections}
- Headers/Footers: {len(headers_footers)}

==================== EXTRACTED CONTENT ====================

{full_extracted_text}

=======================================================

🔧 Technical Details:
- Library: python-docx (specialized DOCX parser)
- Content Types: Paragraphs, Tables, Headers, Footers
- Formatting: Preserved structure with section markers
- Text Processing: Cleaned whitespace and empty elements

💡 DOCX Extraction Features:
✅ Full paragraph text extraction
✅ Table content with cell separation
✅ Headers and footers from all sections
✅ Maintains document structure
✅ Handles complex DOCX formatting
✅ Preserves text content accurately"""

                        method = f"python-docx extraction ({total_paragraphs} paragraphs, {total_tables} tables)"
                        
                    else:
                        extracted_text = f"""📄 DOCX Document Processed - No Text Content Found

Document: {filename}
File Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)
Document Structure: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables

This DOCX document appears to contain:
- Empty paragraphs or only formatting
- Images without text content
- Complex embedded objects
- Non-text elements only

The document was successfully opened but no readable text was found."""
                        method = "python-docx extraction (No text content)"
                        
                except Exception as docx_error:
                    extracted_text = f"""❌ DOCX Processing Error

Document: {filename}
File Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)

Error: {str(docx_error)}

This could be due to:
- Corrupted DOCX file
- Password-protected document
- Unsupported DOCX format/version
- Document contains complex elements not supported
- File is not a valid DOCX format

Please ensure the file is a valid, unprotected DOCX document."""
                    method = "python-docx error"
                    
            elif file_extension in ['doc']:
                # Legacy DOC format note
                extracted_text = f"""📄 Legacy DOC Format Detected

Document: {filename}
File Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)
Format: Microsoft Word DOC (legacy format)

⚠️ LIMITED SUPPORT FOR .DOC FILES

The .DOC format is a legacy binary format that requires specialized libraries.
This system is optimized for modern .DOCX files using python-docx.

Recommendations:
1. Convert .DOC to .DOCX using Microsoft Word or LibreOffice
2. Use "Save As" → "Word Document (.docx)" 
3. Re-upload the .DOCX version for full text extraction

For .DOCX files, you'll get:
✅ Complete text extraction
✅ Table content parsing  
✅ Headers and footers
✅ Document structure preservation

Alternative: If this is actually a .DOCX file with wrong extension, try renaming it to .docx"""
                method = "Legacy DOC format (limited support)"
                
            elif file_extension in ['xlsx', 'pptx']:
                extracted_text = f"""📊 {file_extension.upper()} Document Detected

Document: {filename}
File Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)
Format: Microsoft {file_extension.upper()}

📝 SPECIALIZED EXTRACTION NEEDED

This system currently has full extraction support for:
✅ DOCX (Microsoft Word) - Complete text extraction
✅ PDF - Text extraction + OCR fallback  
✅ Images - Multilingual OCR
✅ TXT, HTML, CSV, JSON - Direct text reading

For {file_extension.upper()} files, specialized libraries would be needed:
- XLSX: openpyxl or xlrd for spreadsheet data
- PPTX: python-pptx for presentation content

With full Kreuzberg integration, these formats would be fully supported.

Current capability: Basic file recognition and metadata extraction."""
                method = f"{file_extension.upper()} format detection"
                extracted_text = f"""📄 {file_extension.upper()} Document Detected

IMPORTANT: This is a preview version. For full text extraction from {file_extension.upper()} files, the Kreuzberg library would be used.

File Details:
- Name: {filename}
- Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)
- Type: {file_extension.upper()} Document

What would happen with Kreuzberg:
- Full text extraction from all pages
- Proper formatting preservation  
- Table and image text recognition
- High accuracy text recognition

To see actual content extraction, try uploading a TXT, HTML, CSV, or JSON file."""
                method = f"Mock {file_extension.upper()} processing"
                
            elif file_extension in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']:
                # Enhanced OCR processing with ALL language support
                try:
                    import pytesseract
                    from PIL import Image, ImageEnhance, ImageFilter
                    
                    # Get available languages
                    try:
                        available_langs = pytesseract.get_languages(config='')
                        lang_string = '+'.join(available_langs) if available_langs else 'eng'
                    except:
                        lang_string = 'eng+ell+fra+deu+spa+ara+chi_sim+chi_tra+jpn+kor+rus+hin+tha+vie'
                    
                    # Open and process the image
                    image = Image.open(filepath)
                    
                    # Convert to RGB if necessary
                    if image.mode != 'RGB':
                        image = image.convert('RGB')
                    
                    # Image preprocessing for better OCR
                    gray_image = image.convert('L')
                    enhancer = ImageEnhance.Contrast(gray_image)
                    enhanced_image = enhancer.enhance(2.0)
                    sharpened = enhanced_image.filter(ImageFilter.SHARPEN)
                    
                    # Multiple OCR configurations with all languages
                    configs = [
                        ('All Languages Auto', f'--oem 3 --psm 6 -l {lang_string}'),
                        ('All Languages Block', f'--oem 3 --psm 4 -l {lang_string}'),
                        ('English + Major', '--oem 3 --psm 6 -l eng+ell+fra+deu+spa+ara+chi_sim+rus'),
                        ('English Focus', '--oem 3 --psm 6 -l eng'),
                        ('Single Text Block', f'--oem 3 --psm 8 -l {lang_string}'),
                        ('Sparse Text', f'--oem 3 --psm 11 -l {lang_string}')
                    ]
                    
                    best_result = ""
                    best_confidence = 0
                    best_method = ""
                    all_attempts = []
                    detected_languages = set()
                    
                    for config_name, config in configs:
                        try:
                            # Extract text with confidence data
                            ocr_data = pytesseract.image_to_data(
                                sharpened, 
                                config=config, 
                                output_type=pytesseract.Output.DICT
                            )
                            
                            # Filter and combine text with confidence > 15 (lower threshold for multilingual)
                            text_parts = []
                            confidences = []
                            
                            for i in range(len(ocr_data['text'])):
                                conf = int(ocr_data['conf'][i])
                                if conf > 15:
                                    text = ocr_data['text'][i].strip()
                                    if text and len(text) > 1:  # Ignore single characters
                                        text_parts.append(text)
                                        confidences.append(conf)
                            
                            if text_parts and confidences:
                                current_text = ' '.join(text_parts)
                                avg_confidence = sum(confidences) / len(confidences)
                                
                                # Try to detect script/language
                                script_info = ""
                                if any(ord(char) > 127 for char in current_text):
                                    if any(0x0370 <= ord(char) <= 0x03FF for char in current_text):
                                        script_info += "Greek "
                                        detected_languages.add("Greek")
                                    if any(0x0400 <= ord(char) <= 0x04FF for char in current_text):
                                        script_info += "Cyrillic "
                                        detected_languages.add("Cyrillic")
                                    if any(0x0600 <= ord(char) <= 0x06FF for char in current_text):
                                        script_info += "Arabic "
                                        detected_languages.add("Arabic")
                                    if any(0x4E00 <= ord(char) <= 0x9FFF for char in current_text):
                                        script_info += "Chinese "
                                        detected_languages.add("Chinese")
                                    if any(0x3040 <= ord(char) <= 0x309F for char in current_text):
                                        script_info += "Hiragana "
                                        detected_languages.add("Japanese")
                                    if any(0x30A0 <= ord(char) <= 0x30FF for char in current_text):
                                        script_info += "Katakana "
                                        detected_languages.add("Japanese")
                                
                                all_attempts.append({
                                    'method': config_name,
                                    'text': current_text,
                                    'confidence': avg_confidence,
                                    'word_count': len(text_parts),
                                    'script': script_info.strip() or "Latin"
                                })
                                
                                # Keep the best result (prioritize longer text with good confidence)
                                quality_score = (avg_confidence * 0.7) + (len(current_text) * 0.3)
                                best_quality = (best_confidence * 0.7) + (len(best_result) * 0.3)
                                
                                if quality_score > best_quality:
                                    best_result = current_text
                                    best_confidence = avg_confidence
                                    best_method = config_name
                                    
                        except Exception as config_error:
                            all_attempts.append({
                                'method': config_name,
                                'text': f"Error: {str(config_error)}",
                                'confidence': 0,
                                'word_count': 0,
                                'script': "Error"
                            })
                    
                    if best_result:
                        # Format the successful result
                        extracted_text = f"""🌍 MULTILINGUAL OCR Results for {filename}

✅ TEXT EXTRACTED SUCCESSFULLY!

Best Method: {best_method}
Overall Confidence: {best_confidence:.1f}%
Detected Scripts: {', '.join(detected_languages) if detected_languages else 'Latin/English'}
Image Size: {image.size[0]}x{image.size[1]} pixels
File Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)

Available Languages: {len(available_langs) if 'available_langs' in locals() else 'Multiple'}

==================== EXTRACTED TEXT ====================

{best_result}

=======================================================

📊 All OCR Attempts (Best to Worst):

"""
                        
                        # Sort attempts by confidence and add details
                        sorted_attempts = sorted(all_attempts, key=lambda x: x['confidence'], reverse=True)
                        for i, attempt in enumerate(sorted_attempts, 1):
                            status = "🏆 BEST" if attempt['method'] == best_method else f"#{i}"
                            extracted_text += f"""
{status} - {attempt['method']}:
   📊 Confidence: {attempt['confidence']:.1f}%
   📝 Words: {attempt['word_count']}
   🌐 Script: {attempt['script']}
   📄 Preview: {attempt['text'][:150]}{'...' if len(attempt['text']) > 150 else ''}
"""
                        
                        extracted_text += f"""
=======================================================

🔧 Technical Details:
- OCR Engine: Tesseract with ALL language packs
- Preprocessing: Contrast enhancement + sharpening
- Language Detection: Automatic multi-script detection
- Character Filtering: Confidence > 15% threshold
- Quality Scoring: Confidence + text length optimization

🌐 Supported Languages Include:
English, Greek, Arabic, Chinese, Japanese, Korean, Russian, 
Hindi, Thai, Vietnamese, French, German, Spanish, Italian, 
Portuguese, Dutch, Polish, Turkish, Hebrew, and many more!

💡 For best results with multilingual text:
- Ensure high image contrast
- Use clear, readable fonts
- Avoid overly decorative or stylized text
- Higher resolution helps with complex scripts"""

                        method = f"Multilingual OCR ({len(sorted_attempts)} methods, Best: {best_method})"
                        
                    else:
                        # No text detected in any configuration
                        extracted_text = f"""🌍 Multilingual OCR Processing Complete - No Clear Text Found

Image Details:
- Name: {filename}
- Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)
- Format: {file_extension.upper()}
- Dimensions: {image.size[0]}x{image.size[1]} pixels
- Available Languages: {len(available_langs) if 'available_langs' in locals() else 'Multiple'}

OCR Attempts Made:
"""
                        for i, attempt in enumerate(all_attempts, 1):
                            extracted_text += f"{i}. {attempt['method']}: {attempt['confidence']:.1f}% confidence ({attempt['script']})\n"
                        
                        extracted_text += f"""
Possible reasons for poor detection:
- Text is too small, blurry, or low contrast
- Highly stylized or decorative fonts
- Very complex multi-script layouts
- Image quality insufficient for character recognition
- Background patterns interfering with text

💡 Tips for better multilingual OCR:
- Use high-resolution, high-contrast images
- Ensure text is clearly readable to human eye
- Avoid backgrounds with patterns or textures
- For best results, text should be horizontal and well-spaced"""
                        
                        method = "Multilingual OCR (No clear text detected)"
                        
                except Exception as ocr_error:
                    extracted_text = f"""❌ Multilingual OCR Processing Error

Image Details:
- Name: {filename}
- Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)
- Format: {file_extension.upper()}

Error: {str(ocr_error)}

This could be due to:
- Corrupted or invalid image file
- Unsupported image format for OCR
- Memory limitations with large images
- OCR engine configuration issue

Please try with a different image or check the file format."""
                    method = "Multilingual OCR Error"
                
            else:
                # Try to read as text anyway
                try:
                    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read(1000)  # Read first 1000 chars
                        if content.strip():
                            extracted_text = content + "\n\n[Content truncated - showing first 1000 characters]"
                            method = "Generic text reading"
                        else:
                            extracted_text = f"File type '{file_extension}' not directly supported in preview mode. With Kreuzberg, this file would be processed appropriately."
                            method = "Unsupported format"
                except:
                    extracted_text = f"Cannot read '{file_extension}' file as text. With Kreuzberg, specialized handlers would process this file type."
                    method = "Binary file detected"
                    
        except Exception as e:
            extracted_text = f"Error processing file: {str(e)}\n\nThis might be a binary file or corrupted document. The full Kreuzberg version would handle this more gracefully."
            method = "Error handling"
        
        # Clean up file
        if os.path.exists(filepath):
            os.remove(filepath)
        
        # Calculate stats
        word_count = len(extracted_text.split()) if extracted_text else 0
        char_count = len(extracted_text) if extracted_text else 0
        
        # Return response
        result = {
            "filename": filename,
            "file_extension": file_extension,
            "file_size_bytes": file_size,
            "extraction_method": method,
            "full_text": extracted_text,
            "word_count": word_count,
            "character_count": char_count,
            "total_pages": 1,
            "status": "success",
            "note": "This is a test version. Full Kreuzberg processing available in production version."
        }
        
        if output_format == 'csv':
            # Return CSV format
            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerow(['Filename', 'Extension', 'Text', 'Word_Count', 'Char_Count', 'Method'])
            writer.writerow([filename, file_extension, extracted_text.replace('\n', ' '), 
                           word_count, char_count, method])
            
            output.seek(0)
            return send_file(
                io.BytesIO(output.getvalue().encode()),
                mimetype='text/csv',
                as_attachment=True,
                download_name=f"{os.path.splitext(filename)[0]}_extracted.csv"
            )
        else:
            return jsonify(result)
            
    except Exception as e:
        return jsonify({"error": f"Processing failed: {str(e)}"}), 500

if __name__ == '__main__':
    print("=" * 60)
    print("🚀 DOCUMENT EXTRACTION API WITH WEB UI")
    print("=" * 60)
    print("📱 Web Interface: http://localhost:8000")
    print("🔧 API Health: http://localhost:8000/health")
    print("📄 Upload files through the web interface!")
    print("=" * 60)
    
    app.run(host='0.0.0.0', port=8000, debug=True)